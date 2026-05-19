import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(report_content: dict, 
                  session_id: str,
                  output_dir: str = "outputs/reports") -> str:
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    
    sections_order = [
        ("title", None),
        ("abstract", "Abstract"),
        ("introduction", "Introduction"),
        ("key_findings", "Key Findings"),
        ("analysis", "Analysis"),
        ("research_gaps", "Research Gaps"),
        ("emerging_trends", "Emerging Trends"),
        ("conclusion", "Conclusion"),
        ("references", "References")
    ]
    
    GREEN = RGBColor(0x15, 0x80, 0x3D)
    DARK = RGBColor(0x1F, 0x29, 0x37)
    
    for key, heading in sections_order:
        content = report_content.get(key, "Content not available")
        if heading is None:
            p = doc.add_heading(str(content), 0)
            p.runs[0].font.color.rgb = GREEN
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            h = doc.add_heading(heading, 1)
            h.runs[0].font.color.rgb = GREEN
            if key == "references" and isinstance(content, list):
                for ref in content:
                    p = doc.add_paragraph(str(ref), style="List Bullet")
                    p.runs[0].font.size = Pt(10)
            else:
                p = doc.add_paragraph(str(content))
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.color.rgb = DARK
    
    path = f"{output_dir}/{session_id}_report.docx"
    doc.save(path)
    return os.path.abspath(path)
